import os
import shutil
import time
from time import sleep
 
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
 
from framework.base.BaseFrame import BaseFrame
from sprider.business.DownLoad import DownLoad
from sprider.business.SeleniumTools import SeleniumTools
from sprider.business.SpriderTools import SpriderTools
 
from selenium import webdriver
from selenium.webdriver.common.by import By
from sprider.model.SpriderEntity import SpriderEntity
from sprider.access.SpriderAccess import SpriderAccess
 
 
class HuaJunCode:
    base_url =  "https://down.chinaz.com" # 采集的网址
    save_path = "D:\\Freedom\\Sprider\\ChinaZ\\"
    sprider_count = 68  # 采集数量
    sprider_start_count=665# 从第几个序号开始 直接改数量即可 会做除法操作正 正在采集第32页的第16个资源 debug
 
 
 
    word_content_list = []
    folder_name = ""
    page_end_number=0
    max_pager=16 #每页的数量
    haved_sprider_count =0  # 已经采集的数量
    page_count = 1  # 每个栏目开始业务content="text/html; charset=gb2312"
    filter_down_file=[]
 
    def __init__(self):
        pass
 
    def sprider(self,title_name="NET"):
 
        """
       采集
       PHP https://down.chinaz.com/class/572_5_1.htm
       NET https://down.chinaz.com/class/572_4_1.htm
       ASP https://down.chinaz.com/class/572_3_1.htm
       Pytyhon https://down.chinaz.com/class/604_572_1.htm
       :return:
       """
        if title_name == "PHP":
            self.folder_name = "PHP源码"
            self.second_column_name = "572_5"
        elif title_name == "Go":
            self.folder_name = "Go源码"
            self.second_column_name = "606_572"
        elif title_name == "NET":
            self.folder_name = "NET源码"
            self.second_column_name = "572_4"
        elif title_name == "ASP":
            self.folder_name = "ASP源码"
            self.second_column_name = "572_3"
        elif title_name == "Python":
            self.folder_name = "Python源码"
            self.second_column_name = "604_572"
        elif title_name == "JavaScript":
            self.folder_name = "JavaScript源码"
            self.second_column_name = "602_572"
        elif title_name == "Java":
            self.folder_name = "Java源码"
            self.second_column_name = "572_517"
        #
 
 
        first_column_name = title_name # 一级目录
        second_folder_name = str(self.sprider_count) + "个" + self.folder_name #二级目录
        self.sprider_type =second_folder_name
        self.merchant=int(self.sprider_start_count) //int(self.max_pager)+1 #起始页码用于效率采集
        self.file_path = self.save_path + os.sep + "Code" + os.sep + first_column_name + os.sep + second_folder_name
        self.save_path = self.save_path+ os.sep + "Code" + os.sep+first_column_name+os.sep + second_folder_name+ os.sep + self.folder_name
        BaseFrame().debug("开始采集ChinaZCode"+self.folder_name+"...")
        sprider_url = (self.base_url + "/class/{0}_1.htm".format(self.second_column_name))
        down_path="D:\\Freedom\\Sprider\\ChinaZ\\Code\\"+first_column_name+"\\"+second_folder_name+"\\Temp\\"
        if os.path.exists(down_path) is True:
            shutil.rmtree(down_path)
        if os.path.exists(down_path) is False:
            os.makedirs(down_path)
 
        if os.path.exists(self.save_path ) is True:
            shutil.rmtree(self.save_path )
        if os.path.exists(self.save_path ) is False:
            os.makedirs(self.save_path )
        chrome_options = webdriver.ChromeOptions()
        diy_prefs ={'profile.default_content_settings.popups': 0,
                    'download.default_directory':'{0}'.format(down_path)}
        # 添加路径到selenium配置中
        chrome_options.add_experimental_option('prefs', diy_prefs)
        chrome_options.add_argument('--headless') #隐藏浏览器
 
        # 实例化chrome浏览器时，关联忽略证书错误
        driver = webdriver.Chrome(options=chrome_options)
        driver.set_window_size(1280, 800)  # 分辨率 1280*800
 
        # driver.get方法将定位在给定的URL的网页，get接受url可以是任何网址，此处以百度为例
        driver.get(sprider_url)
        # content = driver.page_source
        # print(content)
        div_elem = driver.find_element(By.CLASS_NAME, "main")  # 列表页面 核心内容
        element_list = div_elem.find_elements(By.CLASS_NAME, 'item')
 
        laster_pager_ul = driver.find_element(By.CLASS_NAME, "el-pager")
        laster_pager_li =laster_pager_ul.find_elements(By.CLASS_NAME, 'number')
        laster_pager_url = laster_pager_li[len(laster_pager_li) - 1]
        page_end_number = int(laster_pager_url.text)
        self.page_count=self.merchant
        while self.page_count <= int(page_end_number):  # 翻完停止
            try:
                if self.page_count == 1:
                    self.sprider_detail(driver,element_list,self.page_count,page_end_number,down_path)
                    pass
                else:
                    if self.haved_sprider_count == self.sprider_count:
                        BaseFrame().debug("采集到达数量采集停止...")
                        BaseFrame().debug("开始写文章...")
                        self.builder_word(self.folder_name, self.save_path, self.word_content_list)
                        BaseFrame().debug("文件编写完毕，请到对应的磁盘查看word文件和下载文件！")
                        break
                    #(self.base_url + "/sort/{0}/{1}/".format(url_index, self.page_count))
                    #http://soft.onlinedown.net/sort/177/2/
 
                    next_url = self.base_url + "/class/{0}_{1}.htm".format(self.second_column_name, self.page_count)
                    driver.get(next_url)
 
                    div_elem = driver.find_element(By.CLASS_NAME, "main")  # 列表页面 核心内容
                    element_list = div_elem.find_elements(By.CLASS_NAME, 'item')
                    self.sprider_detail( driver, element_list, self.page_count, page_end_number, down_path)
                    pass
                #print(self.page_count)
                self.page_count = self.page_count + 1  # 页码增加1
            except Exception as e:
                print("sprider()执行过程出现错误:" + str(e))
                sleep(1)
 
 
 
    def sprider_detail(self, driver,element_list,page_count,max_page,down_path):
        """
        采集明细页面
        :param driver:
        :param element_list:
        :param page_count:
        :param max_page:
        :param down_path:
        :return:
        """
        index = 0
        element_array=[]
        element_length=len(element_list)
        for element in element_list:
            url_A_obj = element.find_element(By.CLASS_NAME,  'name-text')
            next_url = url_A_obj.get_attribute("href")
            coder_title = url_A_obj.get_attribute("title")
            e=coder_title+"$"+ next_url
            element_array.append(e)
            pass
        if int(self.page_count) == int(self.merchant):
            self.sprider_start_index = int(self.sprider_start_count) % int(self.max_pager)
            index=self.sprider_start_index
        while index < element_length:
 
 
            if os.path.exists(down_path) is False:
                os.makedirs(down_path)
 
            if self.haved_sprider_count == self.sprider_count:
                BaseFrame().debug("采集到达数量采集停止...")
                break
 
            #element = element_list[index]
            element=element_array[index]
            time.sleep(1)
 
            index = index + 1
            sprider_info="正在采集第"+str(page_count)+"页的第"+str(index)+"个资源，共"+str(max_page)+"页资源"
            BaseFrame().debug(sprider_info)
            next_url=element.split("$")[1]
            coder_title=element.split("$")[0]
            # next_url = element.find_element(By.TAG_NAME, 'a').get_attribute("href")
            # coder_title =element.find_element(By.TAG_NAME, 'img').get_attribute("title")
            driver.get(next_url) # 请求明细页面
            try:
                codeEntity = SpriderEntity()  # 下载过的资源不再下载
                codeEntity.sprider_base_url = self.base_url
                codeEntity.create_datetime = SpriderTools.get_current_datetime()
                codeEntity.sprider_url = next_url
                codeEntity.sprider_pic_title = coder_title
                codeEntity.sprider_pic_index = str(index)
                codeEntity.sprider_pager_index = page_count
                codeEntity.sprider_type = self.sprider_type
                if SpriderAccess().query_sprider_entity_by_urlandindex(next_url, str(index)) is None:
                    SpriderAccess().save_sprider(codeEntity)
                else:
                    BaseFrame().debug(coder_title+next_url + "数据采集过因此跳过")
                    continue
 
                if SeleniumTools.judeg_element_isexist(driver, "CLASS_NAME", "download-item") == 3:
                    driver.back()
                    BaseFrame().debug(coder_title+"不存在源码是soft因此跳过哦....")
                    continue
                print("准备点击下载按钮...")
                driver.find_element(By.CLASS_NAME, "download-item").click() #下载源码
 
                result,message=SpriderTools.judge_file_exist(True,240,1,down_path,"zip|rar|gz|tgz")#判断源码
                if result is True:
 
                    sprider_content = [coder_title, self.save_path + os.sep +"image"+ os.sep + coder_title + ".jpg"]  # 采集成功的记录
                    self.word_content_list.append(sprider_content)  # 增加到最终的数组
                    self.haved_sprider_count = self.haved_sprider_count + 1
                    BaseFrame().debug("已经采集完成第" + str(self.haved_sprider_count) + "个")
                    time.sleep(1)
                    driver.back()
 
                    coder_title = str(coder_title).replace("/", "") #去掉windows不识别的字符
                    files = os.listdir(down_path)
                    file_name = files[0] #获取默认值
                    if len(self.filter_down_file)>0:
                        for file in files:
                            for filter_file in self.filter_down_file:
                                if str(file) in str(filter_file):
                                    pass
                                else:
                                    file_name = file
 
                    srcFile = down_path + os.sep + file_name
                    file_ext = os.path.splitext(srcFile)[-1]
 
                    dstFile = down_path + os.sep + coder_title + file_ext
                    os.rename(srcFile, dstFile)
                    srcFile = dstFile
                    dstFile = self.save_path + os.sep + coder_title + file_ext
 
                    shutil.move(srcFile, dstFile)  # 移动文件
 
                else:
                    try:
                        BaseFrame().error("检测下载文件出错可能原因是等待时间不够已经超时，再等待60秒...")
                        time.sleep(60)
                        shutil.rmtree(down_path) #如果没下载完是无法删除的
                        #清空数组
                        self.filter_down_file.clear()
                    except Exception as e:
                        # 使用数组append记录文件名字 移动的时候过滤
                        coder_title = str(coder_title).replace("/", "")  # 去掉windows不识别的字符
                        self.filter_down_file.append(coder_title)
                    pass
            except Exception as e:
                #shutil.rmtree(down_path)
                BaseFrame().error("sprider_detail()执行过程出现错误：" + str(e))
                #driver.get(sprider_url)
                #driver.quit()
 
        if(int(page_count)==int(max_page)):
            self.builder_word(self.folder_name,self.save_path,self.word_content_list)
            BaseFrame().debug("文件编写完毕，请到对应的磁盘查看word文件和下载文件！")
 
    def builder_word(self, word_title, save_path, list_files):
        """
        输出产物是word文件
        :param word_title: 文件的标题
        :param save_path: 文件的保存路径
        :param list_files: 文件集合（单个内容）
        :return:
        """
        try:
            self.copy_file(self.save_path)
 
            print("Create Word"+word_title)
            file_count= len(list_files)
 
            self.gen_passandtxt(file_count,word_title,list_files)
            random_full_file_name = SpriderTools.get_word_image("java",6)
            document = Document()
            document.add_heading(""+word_title+"", level=2)
 
            document.add_paragraph("分享"+str(file_count)+"个"+word_title+"，总有一款适合您\r\n"
                                   "下面是文件的名字，我放了一些图片，文章里不是所有的图主要是放不下...，大家下载后可以看到。")
            document.add_paragraph("源码下载")
            document.add_picture(random_full_file_name, width=Inches(3))
            ppt_tieles = ""
            for files in list_files:
                ppt_tieles = ppt_tieles + str(files[0]) + "\r"
            document.add_paragraph(ppt_tieles)
            # for files in list_files:
            #     try:
            #         document.add_paragraph(files[0])
            #         document.add_picture(files[1], width=Inches(3))
            #     except Exception as e:
            #         pass
            document.add_paragraph("最后送大家一首诗:")
            paragraph = document.add_paragraph()  # 单独控制
            paragraph.add_run("山高路远坑深,\r")
            paragraph.add_run("大军纵横驰奔,\r")
            paragraph.add_run("谁敢横刀立马？\r")
            paragraph.add_run("惟有点赞加关注大军。\r")
            paragraph.bold = True  # 字体加粗
            file_full_path=self.file_path+os.sep+word_title+".docx"
            document.save(file_full_path)
        except Exception as e:
            print("Create Word Fail reason:" + str(e))
    def copy_file(self,target_path):
        print("copy files")
        import os
        import shutil
        src_apk_file_path="薅羊毛专业版.apk"
        dst_apk_file_path=target_path+os.sep+"薅羊毛专业版.apk"
        #shutil.copyfile(src_apk_file_path, dst_apk_file_path)  # 移动文件
 
        src_pdf_file_path = "薅羊毛专业版.pdf"
        dst_pdf_file_path = target_path + os.sep + "薅羊毛专业版.pdf"
        #shutil.copyfile(src_pdf_file_path, dst_pdf_file_path)  # 移动文件
 
        src_pdf_file_path = "亚丁号.url"
        dst_pdf_file_path = self.file_path  + os.sep + "亚丁号.url"
        shutil.copyfile(src_pdf_file_path, dst_pdf_file_path)  # 移动文件
 
        src_doc_file_path = "readme.docx"
        dst_doc_file_path = self.file_path + os.sep + "readme.docx"
        shutil.copyfile(src_doc_file_path, dst_doc_file_path)  # 移动文件
 
        pass
    def gen_passandtxt(self,file_count,word_title, list_files):
 
        print("Create PassWord and Pass.txt")
        message=SpriderTools.gen_password()
        password = "".join(message)
        content=""
        content = content + "\n分享"+str(file_count)+"个"+word_title+"，总有一款适合您"
        content = content + "\n\r"
        content=content+"\n都到这里了您就支持一下呗！谢谢老铁~~"
        content=content+"\n\r"
        content = content + "\n\r"
        content = content + "\n\r"
        for files in list_files:
            content = content+str(files[0])+ "\n"
        content=content+"\n文件我就不一一列举了,送老铁一首打油诗"
        content=content+"\n学习知识费力气，"
        content=content+"\n收集整理更不易。"
        content=content+"\n知识付费甚欢喜，"
        content=content+"\n为咱码农谋福利。"
        content=content+"\n\r"
        content=content+"\n\r"
        content=content+"\n感谢您的支持"
        content=content+"\n\r" 
        content=content+"\n-------------------------------------------华丽分割线-------------------------------------------------------"
        content=content+"\n友情提醒解压密码："+password+""
 
        full_path=self.file_path+os.sep+""+str(file_count)+"sell_pass.txt"
        with open(full_path, 'a', encoding='utf-8') as f:
            f.write(content)
 
 
if __name__ == "__main__":
    HuaJunCode().sprider("Java")
    pass
